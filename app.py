import io
import re
import textwrap
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document

st.set_page_config(
    page_title="ClearView — Business Intelligence",
    page_icon="◈",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ============================================================
# PREMIUM DESIGN — ORIGINAL VISUAL LANGUAGE RETAINED
# ============================================================

st.markdown(textwrap.dedent("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@400;500;600;700&family=Manrope:wght@500;600;700;800&display=swap');
:root{--ink:#0B1220;--muted:#667085;--line:#E7EAF0;--surface:#FFF;--bg:#F5F7FB;--blue:#315CF6;--violet:#6C7CFF;--green:#12B76A;--amber:#F79009;--red:#F04438;--cyan:#06B6D4;--shadow:0 18px 55px rgba(15,23,42,.08)}
*{box-sizing:border-box}html,body,[class*="css"]{font-family:"DM Sans",sans-serif!important}.stApp{background:radial-gradient(circle at 90% 0%,rgba(49,92,246,.09),transparent 26%),radial-gradient(circle at 0% 30%,rgba(108,124,255,.06),transparent 24%),var(--bg)}.block-container{max-width:1420px!important;padding-top:26px!important;padding-bottom:55px!important}#MainMenu,footer,header{visibility:hidden;height:0}
.topbar{display:flex;align-items:center;justify-content:space-between;margin-bottom:22px}.brand{display:flex;align-items:center;gap:11px}.brand-mark{width:40px;height:40px;border-radius:13px;display:flex;align-items:center;justify-content:center;color:#fff;font-family:"Manrope",sans-serif;font-weight:800;background:linear-gradient(135deg,#315CF6,#7C6FF6);box-shadow:0 10px 26px rgba(49,92,246,.25)}.brand-name{font-family:"Manrope",sans-serif;font-size:18px;font-weight:800;color:var(--ink);letter-spacing:-.4px}.brand-pill{padding:8px 13px;border:1px solid var(--line);border-radius:999px;background:rgba(255,255,255,.75);color:#475467;font-size:12px;font-weight:700}
.hero{position:relative;overflow:hidden;min-height:390px;border-radius:32px;padding:54px 58px;color:#fff;background:radial-gradient(circle at 83% 25%,rgba(124,111,246,.40),transparent 22%),radial-gradient(circle at 100% 100%,rgba(49,92,246,.38),transparent 32%),linear-gradient(135deg,#09111F 0%,#111C36 56%,#172A4C 100%);box-shadow:0 26px 75px rgba(11,18,32,.18);animation:fadeUp .6s ease both}.hero:before{content:"";position:absolute;width:520px;height:520px;right:-180px;top:-190px;border:1px solid rgba(255,255,255,.10);border-radius:50%;box-shadow:0 0 0 70px rgba(255,255,255,.025),0 0 0 140px rgba(255,255,255,.018)}.hero-grid{position:relative;z-index:2;display:grid;grid-template-columns:1.35fr .65fr;gap:45px;align-items:center}.eyebrow{display:inline-flex;align-items:center;gap:8px;color:#B9C3FF;font-size:11px;font-weight:800;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:18px}.eyebrow-dot{width:7px;height:7px;border-radius:50%;background:#7C8CFF;box-shadow:0 0 0 6px rgba(124,140,255,.12)}.hero-title{font-family:"Manrope",sans-serif;font-size:clamp(42px,5vw,66px);line-height:1.01;letter-spacing:-3.2px;font-weight:800;margin:0;max-width:820px}.hero-title span{background:linear-gradient(90deg,#A8B5FF,#7CD4FF);-webkit-background-clip:text;background-clip:text;color:transparent}.hero-copy{max-width:760px;color:#C9D1E2;font-size:16px;line-height:1.75;margin-top:22px}.signature{margin-top:25px;color:#AAB3C4;font-size:12px;font-family:"Manrope",sans-serif;letter-spacing:.3px}.signature strong{color:#E9EDFA;font-style:italic;font-size:15px}.hero-visual{min-height:260px;display:flex;align-items:center;justify-content:center}.orbit{position:relative;width:245px;height:245px;border:1px solid rgba(255,255,255,.12);border-radius:50%;animation:spinSlow 18s linear infinite}.orbit:before,.orbit:after{content:"";position:absolute;inset:27px;border:1px solid rgba(255,255,255,.09);border-radius:50%}.orbit:after{inset:59px}.orb-core{position:absolute;inset:82px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-family:"Manrope",sans-serif;font-size:25px;font-weight:800;color:#fff;background:linear-gradient(135deg,#315CF6,#7C6FF6);box-shadow:0 0 50px rgba(108,124,255,.48);animation:pulse 2.8s ease-in-out infinite}.orb-dot{position:absolute;width:12px;height:12px;border-radius:50%;background:#fff;box-shadow:0 0 18px rgba(255,255,255,.75)}.d1{top:16px;left:114px}.d2{right:16px;top:112px}.d3{bottom:22px;left:70px}
.section-head{margin-top:30px;margin-bottom:15px}.section-kicker{color:var(--blue);font-size:11px;font-weight:800;letter-spacing:1.4px;text-transform:uppercase}.section-title{font-family:"Manrope",sans-serif;color:var(--ink);font-size:28px;line-height:1.15;font-weight:800;letter-spacing:-1.1px;margin-top:4px}.section-copy{color:var(--muted);font-size:14px;line-height:1.65;max-width:900px;margin-top:6px}
.workflow{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin:18px 0 28px}.workflow-card{background:rgba(255,255,255,.78);border:1px solid var(--line);border-radius:18px;padding:18px;transition:transform .22s ease,box-shadow .22s ease,border-color .22s ease}.workflow-card:hover{transform:translateY(-4px);box-shadow:var(--shadow);border-color:#D5DAE5}.workflow-no{font-family:"Manrope",sans-serif;font-size:11px;font-weight:800;color:#98A2B3}.workflow-title{font-family:"Manrope",sans-serif;font-size:15px;font-weight:800;color:var(--ink);margin-top:12px}.workflow-text{color:var(--muted);font-size:12px;line-height:1.55;margin-top:5px}
.upload-shell{background:rgba(255,255,255,.88);border:1px solid #DDE2EA;border-radius:25px;padding:25px;box-shadow:0 14px 42px rgba(15,23,42,.06)}.upload-heading{display:flex;align-items:center;justify-content:space-between;gap:15px;margin-bottom:15px}.upload-title{font-family:"Manrope",sans-serif;font-size:20px;font-weight:800;color:var(--ink)}.upload-subtitle{color:var(--muted);font-size:13px;line-height:1.55;margin-top:4px}[data-testid="stFileUploader"]{margin-top:8px}[data-testid="stFileUploaderDropzone"]{background:linear-gradient(180deg,#FBFCFF,#F7F9FD)!important;border:1.5px dashed #AAB4C7!important;border-radius:18px!important;min-height:155px!important}.file-icon{width:30px;height:30px;border-radius:9px;display:flex;align-items:center;justify-content:center;background:#EEF2FF;color:#315CF6;font-weight:800;font-size:12px}
.empty-state{text-align:center;background:rgba(255,255,255,.72);border:1px solid var(--line);border-radius:24px;padding:48px 25px;margin-top:22px;animation:fadeUp .45s ease both}.empty-icon{width:62px;height:62px;border-radius:20px;margin:0 auto 15px;display:flex;align-items:center;justify-content:center;color:#315CF6;font-size:25px;font-weight:800;background:linear-gradient(135deg,#EEF2FF,#F7F4FF)}.empty-title{font-family:"Manrope",sans-serif;font-size:20px;font-weight:800;color:var(--ink)}.empty-copy{color:var(--muted);font-size:13px;line-height:1.6;max-width:620px;margin:7px auto 0}
.metric-card{background:#fff;border:1px solid var(--line);border-radius:18px;padding:18px;min-height:104px;box-shadow:0 8px 25px rgba(15,23,42,.035)}.metric-label{color:#667085;font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.8px}.metric-value{font-family:"Manrope",sans-serif;color:var(--ink);font-size:29px;font-weight:800;letter-spacing:-1px;margin-top:7px}.metric-note{color:#98A2B3;font-size:11px;margin-top:2px}
.insight{background:#fff;border:1px solid var(--line);border-radius:20px;padding:21px;margin:10px 0;box-shadow:0 8px 25px rgba(15,23,42,.035);animation:fadeUp .45s ease both}.insight.problem{border-left:4px solid var(--red)}.insight.solution{border-left:4px solid var(--green)}.insight.suggestion{border-left:4px solid var(--blue)}.insight-label{color:#98A2B3;font-size:10px;font-weight:800;letter-spacing:1.1px;text-transform:uppercase}.insight-title{color:var(--ink);font-family:"Manrope",sans-serif;font-size:18px;font-weight:800;margin-top:6px}.insight-body{color:#475467;font-size:13px;line-height:1.65;margin-top:7px}.priority{display:inline-block;padding:5px 9px;border-radius:999px;font-size:10px;font-weight:800;margin-top:12px;background:#F2F4F7;color:#475467}.priority.high{background:#FEF3F2;color:#B42318}.priority.medium{background:#FFFAEB;color:#B54708}.priority.low{background:#ECFDF3;color:#027A48}
.chart-shell{background:#fff;border:1px solid var(--line);border-radius:22px;padding:18px 20px 10px;box-shadow:0 8px 25px rgba(15,23,42,.035);margin-top:14px}.chart-title{font-family:"Manrope",sans-serif;font-size:17px;font-weight:800;color:var(--ink)}.chart-copy{color:var(--muted);font-size:12px;margin-top:3px}.chart-caption{color:#667085;font-size:12px;margin:8px 0 4px}
.insight-strip{display:grid;grid-template-columns:repeat(3,1fr);gap:12px;margin:14px 0}.decision-card{background:#fff;border:1px solid var(--line);border-radius:18px;padding:18px;box-shadow:0 8px 25px rgba(15,23,42,.035)}.decision-card .tag{font-size:10px;text-transform:uppercase;letter-spacing:1px;font-weight:800}.decision-card h4{font-family:"Manrope",sans-serif;margin:7px 0 4px;color:var(--ink);font-size:16px}.decision-card p{color:#667085;font-size:12px;line-height:1.6;margin:0}.decision-card.red{border-top:3px solid var(--red)}.decision-card.green{border-top:3px solid var(--green)}.decision-card.blue{border-top:3px solid var(--blue)}
.footer{margin-top:45px;padding-top:22px;border-top:1px solid var(--line);text-align:center;color:#98A2B3;font-size:11px;line-height:1.7}.footer strong{color:#667085}.stButton>button{border-radius:11px!important;font-family:"DM Sans",sans-serif!important;font-weight:700!important;min-height:40px!important}.stButton>button:hover{transform:translateY(-1px);box-shadow:0 8px 18px rgba(15,23,42,.08)}
@keyframes fadeUp{from{opacity:0;transform:translateY(10px)}to{opacity:1;transform:translateY(0)}}@keyframes pulse{0%,100%{transform:scale(1);box-shadow:0 0 50px rgba(108,124,255,.42)}50%{transform:scale(1.06);box-shadow:0 0 70px rgba(108,124,255,.60)}}@keyframes spinSlow{from{transform:rotate(0)}to{transform:rotate(360deg)}}
@media(max-width:900px){.hero{padding:38px 28px}.hero-grid{grid-template-columns:1fr}.hero-visual{display:none}.workflow{grid-template-columns:1fr 1fr}.insight-strip{grid-template-columns:1fr}}@media(max-width:600px){.workflow{grid-template-columns:1fr}.hero-title{font-size:42px;letter-spacing:-2px}}
</style>
"""), unsafe_allow_html=True)


def html(markup: str):
    st.markdown(textwrap.dedent(markup), unsafe_allow_html=True)


def safe_text(value):
    return str(value).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def reset_workspace():
    st.session_state.files = {}
    st.session_state.uploader_key += 1


def file_signature(file):
    return f"{file.name}:{file.size}"


def extract_word_document(file):
    document = Document(file)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    tables = []
    for table in document.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append(rows)
    return {"type":"Document","name":file.name,"text":"\n".join(paragraphs),"tables":tables}


def read_file(file):
    name = file.name.lower()
    try:
        if name.endswith(".csv"):
            file.seek(0); return {"type":"Table","name":file.name,"data":pd.read_csv(file)}
        if name.endswith((".xlsx", ".xls")):
            file.seek(0); excel=pd.ExcelFile(file); frames=[]
            for sheet in excel.sheet_names:
                d=pd.read_excel(excel,sheet_name=sheet)
                if not d.empty:
                    d["__source_sheet"]=sheet; frames.append(d)
            return {"type":"Table","name":file.name,"data":pd.concat(frames,ignore_index=True) if frames else pd.DataFrame()}
        if name.endswith(".docx"):
            file.seek(0); return extract_word_document(file)
        if name.endswith(".txt"):
            file.seek(0); raw=file.read(); raw=raw.decode("utf-8",errors="ignore") if isinstance(raw,bytes) else raw
            return {"type":"Document","name":file.name,"text":str(raw),"tables":[]}
        return {"type":"Unsupported","name":file.name,"data":None}
    except Exception as exc:
        return {"type":"Error","name":file.name,"data":str(exc)}


def clean_columns(df):
    d=df.copy(); d.columns=[str(c).strip() for c in d.columns]; return d


def numeric_columns(df): return list(df.select_dtypes(include=np.number).columns)


def normalised_name(name): return re.sub(r"[^a-z0-9]+","_",str(name).lower()).strip("_")


def find_column(df, aliases):
    lookup={normalised_name(c):c for c in df.columns}
    for alias in aliases:
        if normalised_name(alias) in lookup: return lookup[normalised_name(alias)]
    for col in df.columns:
        key=normalised_name(col)
        if any(normalised_name(alias) in key for alias in aliases): return col
    return None


# Broader semantic dictionary. This is the key to making the visual layer adapt to
# real business datasets instead of relying on one fixed demo schema.
SEMANTIC_ALIASES={
    "revenue":["revenue","sales","sales_aed","total_sales","turnover","income","net_sales","gross_revenue"],
    "discount":["discount","discounted","discount_amount","discount_aed"],
    "cost":["cost","operating_cost","expense","expenses","cogs","purchase_cost","total_cost","unit_cost"],
    "profit":["profit","net_profit","gross_profit","ebit","ebitda","margin"],
    "quantity":["quantity","qty","units","units_sold","quantity_in_stock","stock","inventory","order_quantity"],
    "price":["price","unit_price","total_price","selling_price","amount"],
    "customer":["customer","customer_id","client","client_id","buyer","customer_name"],
    "order":["order","order_id","transaction","transaction_id","invoice","invoice_id"],
    "date":["date","order_date","transaction_date","invoice_date","created_at","timestamp","last_order"],
    "region":["region","area","location","territory","branch","store","city"],
    "product":["product","product_id","sku","item","item_name","product_name"],
    "category":["product_category","category","product_type","segment","customer_segment","department"],
    "status":["status","order_status","payment_status","delivery_status"],
    "wait":["waiting_minutes","wait_time","processing_time","lead_time","cycle_time","response_time","working_time","hours"],
    "complaints":["complaints","customer_complaints","returns","refunds","issues","complaint_count"],
    "delay":["delayed_cases","delayed_orders","delay","late_orders","delivery_delay"],
    "stock":["stock","inventory","in_stock","product_in_stock","stock_level"],
    "reorder":["reorder_level","reorder_point","order_level","minimum_stock"],
    "salary":["salary","wage","monthly_salary","annual_salary","pay"],
    "employee":["employee","employee_id","staff","staff_id","worker","worker_id","employee_name"],
    "supplier":["supplier","supplier_name","vendor","vendor_name"],
    "department":["department","team","function","division"],
    "hours":["hours","working_hours","hours_worked","weekly_hours","work_hours"],
}


def infer_semantics(df): return {key:find_column(df,aliases) for key,aliases in SEMANTIC_ALIASES.items()}


def category_options(df):
    s=infer_semantics(df); options=[]
    for key in ["category","product","customer","region","status","department","supplier","employee"]:
        col=s.get(key)
        if col and col not in options and not pd.api.types.is_numeric_dtype(df[col]): options.append(col)
    for col in df.columns:
        if col=="__source_sheet" or col in options or pd.api.types.is_numeric_dtype(df[col]): continue
        n=df[col].nunique(dropna=True)
        if 2<=n<=15: options.append(col)
    return options[:15]


def numeric_metric_options(df):
    s=infer_semantics(df); nums=numeric_columns(df); preferred=[]
    for key in ["revenue","discount","cost","profit","quantity","price","stock","reorder","salary","hours","wait","complaints","delay"]:
        col=s.get(key)
        if col and col in nums and col not in preferred: preferred.append(col)
    return preferred+[c for c in nums if c not in preferred]


def date_column(df):
    s=infer_semantics(df); col=s.get("date")
    if col:
        parsed=pd.to_datetime(df[col],errors="coerce")
        if parsed.notna().sum()>=3: return col
    return None


def human_number(value):
    try:
        value=float(value)
        if abs(value)>=1_000_000:return f"{value/1_000_000:.1f}M"
        if abs(value)>=1_000:return f"{value/1_000:.1f}K"
        return f"{value:,.0f}"
    except Exception:return str(value)


def money_columns(df):
    s=infer_semantics(df); return [c for c in [s.get("revenue"),s.get("discount"),s.get("cost"),s.get("profit"),s.get("salary"),s.get("price")] if c]


def business_context(df):
    s=infer_semantics(df); names=" ".join(normalised_name(c) for c in df.columns)
    if s.get("employee") or s.get("salary") or s.get("hours") or s.get("department") or "staff" in names or "employee" in names:
        return "workforce"
    if s.get("stock") or s.get("reorder") or s.get("supplier") or "inventory" in names or "sku" in names:
        return "inventory"
    if s.get("customer") or s.get("category") or s.get("order"):
        return "sales_customer"
    if s.get("revenue") or s.get("cost") or s.get("profit"):
        return "financial"
    return "general"


# ============================================================
# DECISION ENGINE — BUSINESS PROBLEMS, NOT DATA-CLEANING TASKS
# ============================================================

def analyse_business_table(df):
    df=clean_columns(df); s=infer_semantics(df); ctx=business_context(df)
    problems=[]; solutions=[]; opportunities=[]

    def add_problem(title, evidence, solution, action, priority="medium"):
        problems.append((title,evidence,priority)); solutions.append((title,solution)); opportunities.append((title,action,priority))

    # Revenue leakage: discounts, cost pressure, negative/low margin.
    rev=s.get("revenue"); disc=s.get("discount"); cost=s.get("cost"); profit=s.get("profit")
    if rev:
        r=pd.to_numeric(df[rev],errors="coerce")
        if disc:
            d=pd.to_numeric(df[disc],errors="coerce").fillna(0)
            valid=pd.concat([r.rename("r"),d.rename("d")],axis=1).dropna(subset=["r"])
            total_r=valid.r.sum(); total_d=valid.d.sum()
            if total_r>0 and total_d/total_r>=.05:
                pct=total_d/total_r*100
                add_problem("Revenue leakage from discounting",f"Discounts represent approximately {pct:.1f}% of gross sales in the uploaded data.","Segment discounts by product, customer and order volume; remove blanket discounting where the incremental conversion or retention value does not cover the margin lost.","Prioritise high-discount/low-margin products and customers for discount-rule redesign.","high")
        if cost:
            c=pd.to_numeric(df[cost],errors="coerce"); valid=pd.concat([r.rename("r"),c.rename("c")],axis=1).dropna()
            if not valid.empty and valid.r.sum()>0:
                margin=(valid.r.sum()-valid.c.sum())/valid.r.sum()*100
                if margin<25:
                    add_problem("Margin pressure is limiting profitable growth",f"Calculated gross contribution from the uploaded revenue/cost fields is approximately {margin:.1f}%.","Break the margin down by product, category, supplier and customer segment, then target the largest controllable cost drivers rather than applying a blanket cost cut.","Protect high-contribution products and address low-contribution combinations first.","high")

    if profit:
        p=pd.to_numeric(df[profit],errors="coerce")
        if (p<0).any():
            loss_count=int((p<0).sum())
            add_problem("Loss-making transactions or products are present",f"{loss_count:,} records contain negative profit.","Trace negative-profit records to price, discount, cost and volume drivers; redesign or restrict loss-making combinations where they do not serve a strategic purpose.","Rank loss-making products/customers and quantify their total contribution before deciding on pricing or portfolio changes.","high")

    # Inventory: stockouts / excess stock / reorder risk.
    stock=s.get("stock"); reorder=s.get("reorder"); qty=s.get("quantity")
    if stock and reorder:
        stv=pd.to_numeric(df[stock],errors="coerce"); ro=pd.to_numeric(df[reorder],errors="coerce")
        risk=(stv<ro).sum(); excess=(stv>ro*2).sum()
        if risk:
            add_problem("Stockout risk is above the reorder threshold",f"{int(risk):,} records are below their configured reorder level.","Move from static replenishment to demand-led reorder rules using sales velocity, lead time and supplier reliability.","Prioritise SKUs below reorder level by revenue contribution and lead time.","high")
        if excess:
            add_problem("Excess inventory is tying up working capital",f"{int(excess):,} records hold more than twice their configured reorder level.","Reduce replenishment for slow-moving SKUs and use targeted promotions, bundle offers or supplier renegotiation where appropriate.","Rank excess stock by inventory value and days/velocity so capital is released from the weakest movers.","high")
    elif stock and qty:
        stv=pd.to_numeric(df[stock],errors="coerce")
        if len(stv.dropna())>5 and (stv==0).sum()>0:
            add_problem("Zero-stock items may be causing lost sales",f"{int((stv==0).sum()):,} records show zero stock.","Connect stock levels to sales velocity and lead time to create dynamic reorder points instead of reacting after stock reaches zero.","Prioritise zero-stock products with historical sales or high revenue contribution.","high")

    # Workforce: overstaffing / understaffing / cost per output.
    emp=s.get("employee"); salary=s.get("salary"); hours=s.get("hours"); dept=s.get("department")
    if salary and hours:
        sal=pd.to_numeric(df[salary],errors="coerce"); hrs=pd.to_numeric(df[hours],errors="coerce")
        valid=pd.concat([sal.rename("salary"),hrs.rename("hours")],axis=1).dropna(); valid=valid[valid.hours>0]
        if not valid.empty:
            cost_hour=(valid.salary/valid.hours).replace([np.inf,-np.inf],np.nan).dropna()
            if len(cost_hour)>=5:
                q90=cost_hour.quantile(.9); q10=cost_hour.quantile(.1)
                if q90>q10*1.8:
                    add_problem("Workforce cost is uneven across the available records",f"The highest hourly-equivalent labour costs are more than 1.8× the lower end of the distribution.","Compare pay, scheduled hours, workload and output together before changing staffing. Reallocate hours toward demand peaks and review roles with persistently high cost per unit of output.","Build a cost-versus-hours view by department/role and test schedule changes before permanent headcount decisions.","medium")
    if dept and hours:
        d=pd.DataFrame({"dept":df[dept].astype(str),"hours":pd.to_numeric(df[hours],errors="coerce")}).dropna()
        if not d.empty:
            dep=d.groupby("dept")["hours"].sum().sort_values(ascending=False)
            if len(dep)>=3 and dep.iloc[0]>dep.mean()*1.6:
                add_problem("Workforce hours are concentrated in a small number of teams",f"The highest-hour department accounts for materially more scheduled hours than the average department.","Match labour capacity to measurable workload, customer volume and service targets rather than cutting hours uniformly.","Compare department hours with revenue, orders or workload volume before making schedule or headcount decisions.","medium")

    # Customer / order operations.
    order=s.get("order"); customer=s.get("customer"); status=s.get("status")
    if status and order:
        stc=df[status].fillna("Unknown").astype(str).str.lower(); total=len(stc); bad=(~stc.str.contains("complete|completed|success|paid|delivered",regex=True)).sum()
        if total and bad/total>.15:
            add_problem("Order conversion or fulfilment leakage is material",f"Approximately {bad/total*100:.1f}% of order/status records are outside a completed/successful state.","Break failed orders into payment, stock, fulfilment and customer-initiated reasons, then address the largest failure pathway.","Create a funnel from order created → paid → fulfilled → completed and quantify loss at each step.","high")

    if not problems:
        opportunities.append(("No material business issue triggered by the current rules","The current file does not contain enough evidence for a high-confidence business problem. Add customer, product, cost, inventory or workforce measures to deepen the diagnosis.","low"))
    return problems,solutions,opportunities


def data_quality_notes(df):
    missing=int(df.isna().sum().sum()); dup=int(df.duplicated().sum())
    notes=[]
    if missing: notes.append(f"{missing:,} missing cells may affect confidence in some calculations.")
    if dup: notes.append(f"{dup:,} duplicate rows were detected and should be validated before financial decisions.")
    return notes


def aggregate_decisions(table_results):
    problems=[]; solutions=[]; opportunities=[]
    for r in table_results:
        p,s,o=analyse_business_table(r["data"])
        problems += [(a,b,c,r["name"]) for a,b,c in p]
        solutions += [(a,b,r["name"]) for a,b in s]
        opportunities += [(a,b,c,r["name"]) for a,b,c in o]
    return problems,solutions,opportunities


# ============================================================
# VISUAL ENGINE — MULTIPLE COMPLEMENTARY CHARTS
# ============================================================

def add_labels(fig, template=None):
    fig.update_layout(
        height=430, margin=dict(l=18,r=18,t=58,b=45), paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(family="DM Sans",color="#344054"), hoverlabel=dict(bgcolor="#0B1220",font_color="#FFF"),
        legend=dict(orientation="h",yanchor="bottom",y=1.02,xanchor="right",x=1),
    )
    fig.update_xaxes(showgrid=True,gridcolor="#EEF1F5",zeroline=False,title_font=dict(size=12))
    fig.update_yaxes(showgrid=True,gridcolor="#EEF1F5",zeroline=False,title_font=dict(size=12))
    return fig


def chart_title(title, subtitle):
    html(f'<div class="chart-shell"><div class="chart-title">{safe_text(title)}</div><div class="chart-copy">{safe_text(subtitle)}</div></div>')


def make_visual_suite(df):
    """Return several different visual answers to the same business file."""
    df=clean_columns(df); s=infer_semantics(df); figs=[]
    rev=s.get("revenue"); cost=s.get("cost"); profit=s.get("profit"); disc=s.get("discount"); cat=s.get("category") or s.get("product")
    date=date_column(df); stock=s.get("stock"); reorder=s.get("reorder"); salary=s.get("salary"); hours=s.get("hours"); dept=s.get("department")

    # 1. Revenue vs cost vs profit — grouped bar with different series colors.
    if rev:
        r=pd.to_numeric(df[rev],errors="coerce").fillna(0)
        vals={"Revenue":r.sum()}
        if cost: vals["Cost"]=pd.to_numeric(df[cost],errors="coerce").fillna(0).sum()
        if profit: vals["Profit"]=pd.to_numeric(df[profit],errors="coerce").fillna(0).sum()
        if disc: vals["Discount"]=pd.to_numeric(df[disc],errors="coerce").fillna(0).sum()
        d=pd.DataFrame({"Measure":list(vals),"Amount":list(vals.values())})
        fig=px.bar(d,x="Measure",y="Amount",color="Measure",text_auto=".2s",color_discrete_sequence=["#315CF6","#F04438","#12B76A","#F79009"])
        fig.update_layout(title="Financial position",xaxis_title="Business measure",yaxis_title="Amount")
        figs.append(("Financial position","Revenue, cost, profit and discount are separated so stakeholders can see where value is created or lost.",add_labels(fig)))

    # 2. Category contribution pie/donut.
    if cat and rev:
        d=pd.DataFrame({"Category":df[cat].fillna("Unknown").astype(str),"Revenue":pd.to_numeric(df[rev],errors="coerce")}).dropna()
        d=d.groupby("Category",as_index=False).sum().sort_values("Revenue",ascending=False).head(8)
        if len(d)>=2:
            fig=px.pie(d,names="Category",values="Revenue",hole=.52,color_discrete_sequence=px.colors.qualitative.Set2)
            fig.update_traces(textposition="inside",textinfo="percent+label",hovertemplate="%{label}<br>Revenue: %{value:,.2f}<br>Share: %{percent}<extra></extra>")
            fig.update_layout(title="Revenue mix by category")
            figs.append(("Revenue mix","The donut shows which categories actually carry the revenue base instead of presenting a raw product list.",add_labels(fig)))

    # 3. Category revenue vs cost / profit.
    if cat and rev and cost:
        d=pd.DataFrame({"Category":df[cat].fillna("Unknown").astype(str),"Revenue":pd.to_numeric(df[rev],errors="coerce"),"Cost":pd.to_numeric(df[cost],errors="coerce")}).dropna()
        d=d.groupby("Category",as_index=False).sum().sort_values("Revenue",ascending=False).head(10)
        long=d.melt("Category",var_name="Measure",value_name="Amount")
        fig=px.bar(long,x="Category",y="Amount",color="Measure",barmode="group",color_discrete_map={"Revenue":"#315CF6","Cost":"#F04438"},text_auto=".2s")
        fig.update_layout(title="Revenue versus cost by category",xaxis_tickangle=-25,xaxis_title="Category",yaxis_title="Amount")
        figs.append(("Category economics","This view exposes categories that sell well but consume too much cost, or categories with stronger contribution.",add_labels(fig)))

    # 4. Trend: multiple lines where date exists.
    if date and rev:
        d=pd.DataFrame({"Date":pd.to_datetime(df[date],errors="coerce"),"Revenue":pd.to_numeric(df[rev],errors="coerce")}).dropna()
        if len(d)>=3:
            d=d.groupby("Date",as_index=False).sum().sort_values("Date")
            metrics=["Revenue"]
            if cost:
                d["Cost"]=pd.to_numeric(df.loc[d.index,cost],errors="coerce") if False else pd.to_numeric(df[cost],errors="coerce").groupby(pd.to_datetime(df[date],errors="coerce")).sum().reindex(d["Date"]).values
                metrics.append("Cost")
            fig=go.Figure()
            series_colors={"Revenue":"#315CF6","Cost":"#F04438"}
            for m in metrics:
                fig.add_trace(go.Scatter(x=d["Date"],y=d[m],mode="lines+markers",name=m,line=dict(color=series_colors[m],width=3),marker=dict(size=7)))
            fig.update_layout(title="Performance trend",xaxis_title="Date",yaxis_title="Amount")
            figs.append(("Performance trend","A time-based view highlights growth, deterioration and sudden changes that a category snapshot can hide.",add_labels(fig)))

    # 5. Inventory risk quadrant / scatter.
    if stock and reorder:
        d=pd.DataFrame({"Stock":pd.to_numeric(df[stock],errors="coerce"),"Reorder level":pd.to_numeric(df[reorder],errors="coerce")}).dropna()
        if len(d)>=3:
            d["Status"]=np.where(d["Stock"]<d["Reorder level"],"Below reorder level",np.where(d["Stock"]>2*d["Reorder level"],"Excess stock","Within range"))
            fig=px.scatter(d,x="Reorder level",y="Stock",color="Status",size=np.maximum(d["Stock"].abs(),1),color_discrete_map={"Below reorder level":"#F04438","Excess stock":"#F79009","Within range":"#12B76A"},hover_data={"Reorder level":True,"Stock":True,"Status":True})
            maxv=max(d["Stock"].max(),d["Reorder level"].max())
            fig.add_shape(type="line",x0=0,y0=0,x1=maxv,y1=maxv,line=dict(color="#667085",dash="dash"))
            fig.update_layout(title="Inventory risk map",xaxis_title="Reorder level",yaxis_title="Current stock")
            figs.append(("Inventory risk","Red points indicate potential stockout risk; amber points indicate excess stock and tied-up working capital.",add_labels(fig)))

    # 6. Workforce cost vs hours.
    if salary and hours:
        d=pd.DataFrame({"Salary":pd.to_numeric(df[salary],errors="coerce"),"Hours":pd.to_numeric(df[hours],errors="coerce")}).dropna()
        if len(d)>=3:
            d["Cost per hour"]=d["Salary"]/d["Hours"].replace(0,np.nan)
            if dept: d["Department"]=df.loc[d.index,dept].fillna("Unknown").astype(str)
            else: d["Department"]="Workforce"
            fig=px.scatter(d,x="Hours",y="Salary",color="Department",size="Cost per hour",hover_data=["Cost per hour"],color_discrete_sequence=px.colors.qualitative.Safe)
            fig.update_layout(title="Workforce cost versus working hours",xaxis_title="Working hours",yaxis_title="Salary / labour cost")
            figs.append(("Workforce capacity","The aim is not automatic layoffs: it is to identify where hours, cost and workload do not align so management can test smarter schedules first.",add_labels(fig)))

    # 7. Customer/order status funnel-like bar.
    status=s.get("status")
    if status:
        counts=df[status].fillna("Unknown").astype(str).value_counts().head(10).reset_index(); counts.columns=["Status","Orders"]
        fig=px.bar(counts,x="Status",y="Orders",color="Status",text_auto=True,color_discrete_sequence=px.colors.qualitative.Pastel)
        fig.update_layout(title="Order / service status mix",xaxis_title="Status",yaxis_title="Records")
        figs.append(("Status performance","This makes operational leakage visible: completed, failed, pending and other states can be compared directly.",add_labels(fig)))

    # 8. Discount vs revenue scatter.
    if rev and disc:
        d=pd.DataFrame({"Revenue":pd.to_numeric(df[rev],errors="coerce"),"Discount":pd.to_numeric(df[disc],errors="coerce")}).dropna()
        if len(d)>=3:
            fig=px.scatter(d,x="Revenue",y="Discount",color_discrete_sequence=["#6C7CFF"],trendline="ols" if len(d)>=20 else None)
            fig.update_layout(title="Discounting versus sales",xaxis_title="Revenue / sales",yaxis_title="Discount")
            figs.append(("Discount leakage","High discounts attached to low-value transactions are a direct place to investigate margin leakage and discount-rule effectiveness.",add_labels(fig)))

    return figs


# ============================================================
# BRAND + LANDING
# ============================================================
html("""
<div class="topbar"><div class="brand"><div class="brand-mark">C</div><div class="brand-name">ClearView</div></div><div class="brand-pill">Business Intelligence</div></div>
<div class="hero"><div class="hero-grid"><div><div class="eyebrow"><span class="eyebrow-dot"></span> Business Operations Intelligence</div><div class="hero-title">See what is happening.<br><span>Know what to do next.</span></div><div class="hero-copy">ClearView turns the business information you already have into one clear, interactive decision view — showing what needs attention, why it matters, what can be done, and where the next opportunity may be.</div><div class="signature">ATW · independent business intelligence project</div></div><div class="hero-visual"><div class="orbit"><div class="orb-core">CV</div><div class="orb-dot d1"></div><div class="orb-dot d2"></div><div class="orb-dot d3"></div></div></div></div></div>
""")
html("""
<div class="section-head"><div class="section-kicker">How ClearView works</div><div class="section-title">One workspace. From business files to decisions.</div><div class="section-copy">Upload spreadsheets, CSV exports and Word documents. ClearView identifies business signals, diagnoses likely operational problems, and presents the result through decision-focused visualisations rather than a raw list of fields.</div></div>
<div class="workflow"><div class="workflow-card"><div class="workflow-no">01</div><div class="workflow-title">Upload</div><div class="workflow-text">Add multiple files from the same business.</div></div><div class="workflow-card"><div class="workflow-no">02</div><div class="workflow-title">Understand</div><div class="workflow-text">Identify revenue, costs, customers, products, inventory and workforce signals.</div></div><div class="workflow-card"><div class="workflow-no">03</div><div class="workflow-title">Diagnose</div><div class="workflow-text">Translate patterns into business problems such as leakage, stock risk, margin pressure or capacity imbalance.</div></div><div class="workflow-card"><div class="workflow-no">04</div><div class="workflow-title">Act</div><div class="workflow-text">Show a practical intervention and the business outcome it is intended to improve.</div></div></div>
""")

html('<div class="upload-shell"><div class="upload-heading"><div><div class="upload-title">Upload your business information</div><div class="upload-subtitle">CSV · XLSX · XLS · DOCX · TXT · multiple files supported</div></div></div>')

if "files" not in st.session_state: st.session_state.files={}
if "uploader_key" not in st.session_state: st.session_state.uploader_key=0

uploaded=st.file_uploader("Choose one or more business files",type=["csv","xlsx","xls","docx","txt"],accept_multiple_files=True,key=f"uploader_{st.session_state.uploader_key}")
if uploaded:
    for f in uploaded: st.session_state.files[file_signature(f)]=f

if st.session_state.files:
    st.markdown("**Files in this workspace**")
    for key in list(st.session_state.files):
        f=st.session_state.files[key]; c1,c2,c3=st.columns([.08,.72,.20]); suffix=Path(f.name).suffix.upper().replace(".","") or "FILE"
        with c1: st.markdown(f"<div class='file-icon'>{safe_text(suffix[:4])}</div>",unsafe_allow_html=True)
        with c2: st.markdown(f"**{safe_text(f.name)}**<br><span style='color:#98A2B3;font-size:11px'>{f.size/1024:.0f} KB</span>",unsafe_allow_html=True)
        with c3:
            if st.button("Remove",key=f"remove_{key}",use_container_width=True): del st.session_state.files[key]; st.rerun()
    if st.button("Clear all files",key="clear_all"):
        reset_workspace(); st.rerun()
st.markdown('</div>',unsafe_allow_html=True)

if not st.session_state.files:
    html('<div class="empty-state"><div class="empty-icon">◈</div><div class="empty-title">Your business workspace is ready.</div><div class="empty-copy">Add your files above and ClearView will build the decision view from those files.</div></div><div class="footer">ClearView — Business Intelligence<br><strong>ATW</strong> · an independent business analytics project<br>© 2026 Anjalo Theophine Wilson · All rights reserved</div>')
    st.stop()

results=[read_file(f) for f in st.session_state.files.values()]
table_results=[r for r in results if r["type"]=="Table"]
document_results=[r for r in results if r["type"]=="Document"]
errors=[r for r in results if r["type"] in ("Error","Unsupported")]
for r in table_results: r["data"]=clean_columns(r["data"])

total_rows=sum(len(r["data"]) for r in table_results); total_columns=sum(len(r["data"].columns) for r in table_results); total_missing=sum(int(r["data"].isna().sum().sum()) for r in table_results)

html('<div class="section-head"><div class="section-kicker">Executive snapshot</div><div class="section-title">The business at a glance.</div><div class="section-copy">These headline measures show the size of the evidence base before the diagnostic view.</div></div>')
m1,m2,m3,m4=st.columns(4)
for col,(label,value,note) in zip((m1,m2,m3,m4),[("Files",len(results),"in this workspace"),("Records",f"{total_rows:,}","rows across tables"),("Fields",f"{total_columns:,}","columns across tables"),("Documents",len(document_results),"Word / text files")]):
    with col: html(f'<div class="metric-card"><div class="metric-label">{safe_text(label)}</div><div class="metric-value">{safe_text(value)}</div><div class="metric-note">{safe_text(note)}</div></div>')

# ============================================================
# DECISION VIEW
# ============================================================
problems,solutions,opportunities=aggregate_decisions(table_results)
html('<div class="section-head"><div class="section-kicker">Decision view</div><div class="section-title">Problem → Business impact → Action.</div><div class="section-copy">ClearView deliberately avoids generic recommendations such as “review the data”. Each triggered issue is translated into a business problem, an intervention and the management decision it supports.</div></div>')

if problems:
    for title,evidence,priority,source in problems[:8]:
        matching=[x for x in solutions if x[0]==title and x[2]==source]
        action=[x for x in opportunities if x[0]==title and x[3]==source]
        sol=matching[0][1] if matching else "Investigate the driver and test a targeted operational intervention."
        act=action[0][1] if action else "Prioritise the issue using measurable business impact."
        html(f'<div class="insight-strip"><div class="decision-card red"><div class="tag" style="color:#F04438">Problem · {safe_text(source)}</div><h4>{safe_text(title)}</h4><p>{safe_text(evidence)}</p></div><div class="decision-card green"><div class="tag" style="color:#12B76A">Solution</div><h4>Recommended intervention</h4><p>{safe_text(sol)}</p></div><div class="decision-card blue"><div class="tag" style="color:#315CF6">Management action</div><h4>What to do next</h4><p>{safe_text(act)}</p></div></div>')
else:
    html('<div class="insight solution"><div class="insight-label">No high-confidence problem triggered</div><div class="insight-title">The current file does not support a specific business diagnosis.</div><div class="insight-body">Add more relevant business measures such as product/category, customer, revenue, cost, stock, supplier, order status, employee, salary or working-hours data.</div></div>')

# ============================================================
# MULTI-GRAPH STAKEHOLDER DASHBOARD
# ============================================================
for idx,result in enumerate(table_results):
    df=result["data"]; ctx=business_context(df)
    html(f'<div class="section-head"><div class="section-kicker">Stakeholder visual dashboard · {safe_text(result["name"])}</div><div class="section-title">See the story, not just the rows.</div><div class="section-copy">Different charts answer different business questions: contribution, mix, trend, leakage, inventory risk, order performance and workforce capacity.</div></div>')
    figs=make_visual_suite(df)
    if not figs:
        st.info("The file does not contain enough recognised business fields for the specialised visual dashboard. The raw data remains available below.")
    else:
        for j,(title,desc,fig) in enumerate(figs):
            chart_title(title,desc)
            st.plotly_chart(fig,use_container_width=True,config={"displaylogo":False,"responsive":True})
    with st.expander(f"Inspect {result['name']} data",expanded=False):
        st.dataframe(df,use_container_width=True,height=330)

    # Compact business interpretation based on detected domain.
    if ctx=="workforce":
        html('<div class="insight suggestion"><div class="insight-label">Workforce interpretation</div><div class="insight-title">Use cost + hours + workload together before changing headcount.</div><div class="insight-body">A dashboard should identify under-utilised capacity, overtime pressure and cost concentration. It should not automatically tell management to dismiss employees. The better analytical sequence is demand → workload → required capacity → schedule optimisation → only then structural headcount decisions if supported by evidence.</div></div>')
    elif ctx=="inventory":
        html('<div class="insight suggestion"><div class="insight-label">Inventory interpretation</div><div class="insight-title">Separate stockout risk from excess-stock risk.</div><div class="insight-body">A single “inventory” number hides two opposite problems. ClearView therefore maps current stock against reorder level so stakeholders can see what may cause lost sales and what may be tying up cash.</div></div>')
    elif ctx=="sales_customer":
        html('<div class="insight suggestion"><div class="insight-label">Customer interpretation</div><div class="insight-title">Connect customer behaviour to commercial outcomes.</div><div class="insight-body">Use category, customer, order status, discount and revenue relationships to identify where conversion, retention or margin is being lost rather than treating every transaction equally.</div></div>')

# ============================================================
# DATA QUALITY — SUPPORTING INFORMATION, NOT THE MAIN DIAGNOSIS
# ============================================================
if table_results:
    quality=[]
    for r in table_results:
        q=data_quality_notes(r["data"])
        if q: quality.append((r["name"],q))
    if quality:
        html('<div class="section-head"><div class="section-kicker">Evidence confidence</div><div class="section-title">Data-quality checks that may affect decisions.</div><div class="section-copy">These checks are intentionally separated from the business diagnosis. A missing value is a data issue; it is not automatically a business problem.</div></div>')
        for name,notes in quality:
            for note in notes: html(f'<div class="insight suggestion"><div class="insight-label">Data confidence · {safe_text(name)}</div><div class="insight-body">{safe_text(note)}</div></div>')

# ============================================================
# DOCUMENT EVIDENCE
# ============================================================
if document_results:
    html('<div class="section-head"><div class="section-kicker">Supporting evidence</div><div class="section-title">Documents are part of the business story.</div><div class="section-copy">Word and text files are extracted and surfaced instead of being treated as irrelevant uploads.</div></div>')
    for doc in document_results:
        words=len(re.findall(r"\b\w+\b",doc.get("text",""))); tables=len(doc.get("tables",[]))
        html(f'<div class="insight suggestion"><div class="insight-label">Document reviewed</div><div class="insight-title">{safe_text(doc["name"])}</div><div class="insight-body">ClearView extracted approximately <strong>{words:,}</strong> words and <strong>{tables:,}</strong> table(s).</div></div>')
        with st.expander(f"Read extracted content · {doc['name']}"):
            st.text_area("Extracted text",doc.get("text","")[:30000],height=260,key=f"text_{doc['name']}")
            for ti,table in enumerate(doc.get("tables",[])):
                st.markdown("**Table found in document**"); st.dataframe(pd.DataFrame(table),use_container_width=True,height=220,key=f"doc_table_{doc['name']}_{ti}")

if errors:
    html('<div class="insight problem"><div class="insight-label">File handling</div><div class="insight-title">Some files could not be analysed.</div><div class="insight-body">Remove the affected file and replace it with a readable CSV, Excel, Word or text file.</div></div>')
    for error in errors: st.caption(f"{error['name']}: {error.get('data','Unsupported file type')}")

html('<div class="footer">ClearView — Business Intelligence<br><strong>ATW</strong> · an independent business analytics project<br>© 2026 Anjalo Theophine Wilson · All rights reserved</div>')